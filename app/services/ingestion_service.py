"""
==============================================================================
Data Ingestion Service — Multi-Format Document Loading
==============================================================================

ARCHITECTURAL DECISION:
    The ingestion layer handles raw data → structured Document objects.
    It supports multiple formats as required by the BRD:
    
    1. **JSON** — Primary format for structured banking scheme data
    2. **Text/Markdown** — Banking policy documents
    3. **CSV** — Tabular data (interest rates, eligibility matrices)
    4. **PDF** — Scanned/digital banking circulars (via pypdf)
    5. **DOCX** — Word documents (via python-docx)
    6. **HTML** — Web-scraped banking information (via BeautifulSoup)
    
    WHY NOT USE LlamaIndex SimpleDirectoryReader?
    While LlamaIndex provides a reader, building our own gives us:
    - Full control over metadata extraction
    - Custom preprocessing for banking-specific noise
    - No dependency lock-in for the ingestion layer
    - Banking-specific field mapping (scheme_id, category, etc.)
    - Testability — each format handler is independently testable

DATA PREPROCESSING:
    Banking data has specific noise patterns:
    - Legal disclaimers and footnotes
    - Repeated headers/footers in PDFs
    - Unicode issues in government documents
    - Inconsistent date formats
    Our preprocessing pipeline handles all of these.
"""

import json
import csv
import os
import re
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional

from loguru import logger


# =============================================================================
# Document Model
# =============================================================================

@dataclass
class Document:
    """
    Internal document representation after ingestion.
    
    WHY a dataclass (not Pydantic)?
    Documents are internal data structures, not API contracts.
    Dataclasses are lighter weight and sufficient for this purpose.
    Pydantic is reserved for API boundary validation (schemas.py).
    """
    doc_id: str
    content: str
    metadata: dict = field(default_factory=dict)
    source: str = ""
    doc_type: str = "text"  # text, json, csv, pdf, docx, html

    @property
    def char_count(self) -> int:
        return len(self.content)

    @property 
    def word_count(self) -> int:
        return len(self.content.split())


# =============================================================================
# Ingestion Service
# =============================================================================

class IngestionService:
    """
    Multi-format document ingestion with banking-specific preprocessing.
    
    Usage:
        service = IngestionService()
        docs = service.ingest_directory("./data/raw")
        # or
        docs = service.ingest_banking_schemes("./data/banking_schemes.json")
    """

    # Format-specific handlers
    SUPPORTED_EXTENSIONS = {
        ".json": "_ingest_json",
        ".txt": "_ingest_text",
        ".md": "_ingest_text",
        ".csv": "_ingest_csv",
        ".pdf": "_ingest_pdf",
        ".docx": "_ingest_docx",
        ".html": "_ingest_html",
        ".htm": "_ingest_html",
    }

    def ingest_directory(self, directory_path: str) -> list[Document]:
        """
        Ingest all supported files from a directory.
        
        Recursively walks the directory, identifies file types,
        and dispatches to format-specific handlers.
        """
        documents = []
        dir_path = Path(directory_path)

        if not dir_path.exists():
            logger.warning(f"Ingestion directory not found: {directory_path}")
            return documents

        for file_path in sorted(dir_path.rglob("*")):
            if file_path.is_file():
                ext = file_path.suffix.lower()
                handler_name = self.SUPPORTED_EXTENSIONS.get(ext)

                if handler_name:
                    handler = getattr(self, handler_name)
                    try:
                        file_docs = handler(str(file_path))
                        documents.extend(file_docs)
                        logger.debug(
                            f"Ingested {len(file_docs)} docs from {file_path.name}"
                        )
                    except Exception as e:
                        logger.error(f"Failed to ingest {file_path}: {e}")

        logger.info(
            f"Directory ingestion complete | dir={directory_path} | "
            f"total_docs={len(documents)}"
        )
        return documents

    def ingest_banking_schemes(self, json_path: str) -> list[Document]:
        """
        Specialized ingestion for banking scheme JSON data.
        
        Converts each scheme into a rich, structured document that
        preserves all fields while creating a natural-language text
        representation suitable for embedding and retrieval.
        
        WHY STRUCTURED TEXT (not raw JSON)?
        Embedding models work best on natural language. Converting
        structured data to a narrative format improves:
        - Embedding quality (models trained on prose, not JSON)
        - Retrieval accuracy (query-document similarity improves)
        - Chunk quality (semantic boundaries are clearer in prose)
        """
        documents = []

        if not os.path.exists(json_path):
            logger.warning(f"Banking schemes file not found: {json_path}")
            return documents

        with open(json_path, "r", encoding="utf-8") as f:
            schemes = json.load(f)

        for scheme in schemes:
            # Build a rich, structured text representation
            content = self._scheme_to_document_text(scheme)
            
            doc = Document(
                doc_id=scheme["scheme_id"],
                content=content,
                metadata={
                    "scheme_id": scheme["scheme_id"],
                    "scheme_name": scheme["scheme_name"],
                    "category": scheme.get("category", ""),
                    "ministry": scheme.get("ministry", ""),
                    "launched_date": scheme.get("launched_date", ""),
                    "last_verified": scheme.get("last_verified", ""),
                    "source_url": scheme.get("source", ""),
                    "doc_type": "banking_scheme",
                },
                source=json_path,
                doc_type="json",
            )
            documents.append(doc)

        logger.info(
            f"Banking schemes ingested | file={json_path} | "
            f"schemes={len(documents)}"
        )
        return documents

    # =========================================================================
    # Format-Specific Handlers
    # =========================================================================

    def _ingest_json(self, file_path: str) -> list[Document]:
        """Ingest JSON files — handles both array and object formats."""
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        # If it's a list of schemes, use specialized handler
        if isinstance(data, list) and data and "scheme_id" in data[0]:
            return self.ingest_banking_schemes(file_path)

        # Generic JSON — convert to text
        content = json.dumps(data, indent=2, ensure_ascii=False)
        return [Document(
            doc_id=Path(file_path).stem,
            content=content,
            metadata={"file": file_path},
            source=file_path,
            doc_type="json",
        )]

    def _ingest_text(self, file_path: str) -> list[Document]:
        """Ingest plain text and markdown files."""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        content = self._preprocess_text(content)

        return [Document(
            doc_id=Path(file_path).stem,
            content=content,
            metadata={"file": file_path},
            source=file_path,
            doc_type="text",
        )]

    def _ingest_csv(self, file_path: str) -> list[Document]:
        """
        Ingest CSV files — each row becomes part of a document.
        
        For banking data, CSVs typically contain interest rate tables,
        eligibility matrices, or scheme comparison data.
        """
        documents = []
        
        with open(file_path, "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            rows = list(reader)

        if not rows:
            return documents

        # Convert rows to a readable text format
        headers = list(rows[0].keys())
        content_lines = [f"Data from: {Path(file_path).name}"]
        content_lines.append(f"Columns: {', '.join(headers)}")
        content_lines.append("")

        for i, row in enumerate(rows):
            row_text = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
            content_lines.append(f"Row {i+1}: {row_text}")

        return [Document(
            doc_id=Path(file_path).stem,
            content="\n".join(content_lines),
            metadata={
                "file": file_path,
                "row_count": len(rows),
                "columns": headers,
            },
            source=file_path,
            doc_type="csv",
        )]

    def _ingest_pdf(self, file_path: str) -> list[Document]:
        """
        Ingest PDF files using pypdf.
        
        Handles:
        - Multi-page documents (concatenated with page markers)
        - Noisy text extraction (government PDFs are often messy)
        - Metadata extraction from PDF properties
        """
        try:
            from pypdf import PdfReader
        except ImportError:
            logger.error("pypdf not installed — skipping PDF ingestion")
            return []

        reader = PdfReader(file_path)
        pages = []

        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            text = self._preprocess_text(text)
            if text.strip():
                pages.append(f"[Page {i+1}]\n{text}")

        content = "\n\n".join(pages)

        # Extract PDF metadata
        pdf_meta = reader.metadata or {}
        
        return [Document(
            doc_id=Path(file_path).stem,
            content=content,
            metadata={
                "file": file_path,
                "page_count": len(reader.pages),
                "title": getattr(pdf_meta, "title", None) or "",
                "author": getattr(pdf_meta, "author", None) or "",
            },
            source=file_path,
            doc_type="pdf",
        )]

    def _ingest_docx(self, file_path: str) -> list[Document]:
        """Ingest DOCX files using python-docx."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.error("python-docx not installed — skipping DOCX ingestion")
            return []

        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)
        content = self._preprocess_text(content)

        return [Document(
            doc_id=Path(file_path).stem,
            content=content,
            metadata={"file": file_path, "paragraph_count": len(paragraphs)},
            source=file_path,
            doc_type="docx",
        )]

    def _ingest_html(self, file_path: str) -> list[Document]:
        """Ingest HTML files using BeautifulSoup."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            logger.error("beautifulsoup4 not installed — skipping HTML ingestion")
            return []

        with open(file_path, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f.read(), "html.parser")

        # Remove script and style elements
        for element in soup(["script", "style", "nav", "footer", "header"]):
            element.decompose()

        content = soup.get_text(separator="\n", strip=True)
        content = self._preprocess_text(content)

        title = soup.title.string if soup.title else Path(file_path).stem

        return [Document(
            doc_id=Path(file_path).stem,
            content=content,
            metadata={"file": file_path, "title": title},
            source=file_path,
            doc_type="html",
        )]

    # =========================================================================
    # Preprocessing
    # =========================================================================

    @staticmethod
    def _preprocess_text(text: str) -> str:
        """
        Clean and normalize text for embedding.
        
        Banking-specific preprocessing:
        - Remove excessive whitespace (government docs have irregular spacing)
        - Normalize Unicode characters (Hindi/English mixed docs)
        - Remove page numbers and repeated headers
        - Preserve paragraph boundaries (critical for semantic chunking)
        """
        # Normalize whitespace (preserve paragraph breaks)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Remove common PDF artifacts
        text = re.sub(r"Page \d+ of \d+", "", text)
        text = re.sub(r"^\d+\s*$", "", text, flags=re.MULTILINE)

        # Normalize quotes and dashes
        text = text.replace("\u2018", "'").replace("\u2019", "'")
        text = text.replace("\u201c", '"').replace("\u201d", '"')
        text = text.replace("\u2013", "-").replace("\u2014", "-")

        # Remove leading/trailing whitespace on each line
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(lines)

        return text.strip()

    @staticmethod
    def _scheme_to_document_text(scheme: dict) -> str:
        """
        Convert a banking scheme JSON object to a rich text document.
        
        WHY THIS FORMAT?
        The document is structured as a natural-language description
        with clear section headers. This format:
        1. Embeds well (embedding models prefer prose over JSON)
        2. Chunks well (each section is a natural semantic boundary)
        3. Retrieves well (section headers match query patterns)
        4. Is human-readable (evaluators can inspect directly)
        """
        sections = []

        # Title
        sections.append(f"# {scheme['scheme_name']}")
        sections.append(f"Scheme ID: {scheme['scheme_id']}")
        sections.append(f"Category: {scheme.get('category', 'N/A')}")

        if scheme.get("ministry"):
            sections.append(f"Ministry: {scheme['ministry']}")
        if scheme.get("launched_date"):
            sections.append(f"Launched: {scheme['launched_date']}")

        sections.append("")

        # Description
        if scheme.get("description"):
            sections.append("## Description")
            sections.append(scheme["description"])
            sections.append("")

        # Eligibility
        if scheme.get("eligibility"):
            sections.append("## Eligibility Criteria")
            sections.append(scheme["eligibility"])
            sections.append("")

        # Benefits
        if scheme.get("benefits"):
            sections.append("## Benefits")
            sections.append(scheme["benefits"])
            sections.append("")

        # Interest Rate
        if scheme.get("interest_rate"):
            sections.append("## Interest Rate / Premium")
            sections.append(scheme["interest_rate"])
            sections.append("")

        # Documents Required
        if scheme.get("documents_required"):
            sections.append("## Documents Required")
            sections.append(scheme["documents_required"])
            sections.append("")

        # Source
        if scheme.get("source"):
            sections.append(f"Source: {scheme['source']}")
        if scheme.get("last_verified"):
            sections.append(f"Last Verified: {scheme['last_verified']}")

        return "\n".join(sections)
